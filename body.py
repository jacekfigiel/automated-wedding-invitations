from docx import Document

doc = Document("Invitation.docx")
inv_docx = Document("List of guest.docx")
text = doc.paragraphs[1]

all_guests = []

all_paragraphs = inv_docx.paragraphs
for para in all_paragraphs:
    all_guests.append(para.text)

for g in all_guests:
    prior_paragraph = text.insert_paragraph_before(text=g,
                                                   style="Normal")
    file_name = "Invitation" + g
    doc.save("Invitation for" + g + ".docx")